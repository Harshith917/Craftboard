# -*- coding: utf-8 -*-
"""Generate the Craftboard project overview report (.docx and .html)."""
import datetime
import html
import os

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(HERE, "Craftboard_Project_Overview.docx")
HTML_PATH = os.path.join(HERE, "Craftboard_Project_Overview.html")

TITLE = "Craftboard"
SUBTITLE = "A Real-Time Collaborative Design Canvas"
DOC_TITLE = "Project Overview & Technical Documentation"
VERSION = "1.0"
TODAY = datetime.date.today().strftime("%B %d, %Y")

PALETTE = {
    "ink": "#0f172a",
    "slate": "#475569",
    "sky": "#0ea5e9",
    "sky_dark": "#0369a1",
    "line": "#cbd5e1",
    "bg": "#f8fafc",
    "white": "#ffffff",
}

CONTENT = {
    "sections": [
        {
            "num": "1",
            "title": "Executive Summary",
            "paras": [
                "Craftboard is a web-based, real-time collaborative design canvas that lets distributed teams "
                "ideate and design together on a shared, structured whiteboard. It combines a Konva-powered canvas "
                "frontend with a PostgreSQL-backed Express API, Liveblocks for real-time synchronization and presence, "
                "Clerk for authentication, and a locally hosted Ollama AI assistant that turns natural-language prompts "
                "into ready-to-place design layouts.",
                "The product is designed as a lightweight, self-hostable alternative to heavyweight proprietary design "
                "tools. Because the AI inference runs on a local Ollama instance rather than a third-party cloud API, "
                "teams keep full control over their design data, avoid per-seat SaaS costs, and retain offline-capable "
                "operation.",
                "This document summarizes what was built, how the system works, the technology choices behind it, the "
                "verification performed, and the recommended path forward.",
            ],
        },
        {
            "num": "2",
            "title": "Problem Statement",
            "paras": [
                "Modern product and design teams frequently struggle with three related challenges:",
            ],
            "bullets": [
                "Collaboration friction: sharing designs through static screenshots and file email chains creates "
                "version confusion and a lack of a single source of truth.",
                "Heavyweight tooling: mainstream design platforms are either proprietary SaaS with rising per-seat "
                "pricing, or overly complex self-hosted stacks that are expensive to operate.",
                "Privacy and cost of AI: AI-assisted design is usually gated behind cloud APIs that raise data-privacy "
                "concerns, add latency, and incur usage fees.",
            ],
            "paras2": [
                "Craftboard addresses all three by providing a real-time, role-based collaborative canvas with a "
                "structured shape model, self-hosted infrastructure, and a local AI design assistant.",
            ],
        },
        {
            "num": "3",
            "title": "Solution Overview",
            "paras": [
                "Craftboard organizes work into projects, each containing one or more pages (canvases). Users sign in "
                "through Clerk, are grouped into teams, and are assigned roles that control what they may do on a "
                "project (owner, editor, viewer). Any change to a canvas is synchronized to other connected clients in "
                "real time through Liveblocks, and asynchronously persisted to PostgreSQL.",
                "A built-in access-management workflow lets users request access to a project and lets the owner or an "
                "editor approve or deny those requests from the dashboard. A Socket.IO layer delivers real-time "
                "notifications for these events.",
                "The AI assistant, served by an Express route and powered by a local Ollama deployment, accepts a "
                "plain-language instruction such as \u201cbuild a hero section with a navy background and a call-to-action "
                "button\u201d and returns a structured set of shapes that the client renders directly onto the canvas.",
            ],
        },
        {
            "num": "4",
            "title": "How It Works",
            "subsections": [
                {
                    "title": "4.1 User Journey",
                    "paras": [
                        "A visitor lands on the public marketing site, reviews the feature set and architecture "
                        "diagram, and signs up with email, Google, or GitHub via Clerk. After sign-in the user is "
                        "redirected to the dashboard, which lists projects, recent activity, and pending access "
                        "requests.",
                        "The user creates a project, then a page inside it, and opens the editor. On the canvas they "
                        "can add and edit over a dozen structured shape types, move and resize them, and delete them. "
                        "If they have been granted edit access, they can invite collaborators by email or link.",
                        "To bootstrap a layout, the user opens the AI assistant panel, enters a description, and "
                        "chooses to have the generated shapes added to the canvas for refinement. When the user leaves "
                        "the page, all changes are persisted to the database.",
                    ],
                },
                {
                    "title": "4.2 Real-Time Synchronization Model",
                    "paras": [
                        "Each page is backed by a dedicated Liveblocks room. Connected clients keep a local copy of "
                        "the room state, apply optimistic updates instantly, and reconcile with peers through "
                        "Liveblocks' conflict-free data structures. Live cursor positions and presence data provide "
                        "the live-collaboration feel.",
                        "Persistence is intentionally decoupled from the real-time path. A Liveblocks "
                        "storageUpdated webhook is delivered to the Express API and written to PostgreSQL, so the "
                        "database is never on the critical path of an interactive edit. The client also fetches room "
                        "initial state on entry so a user joining a session receives the current design immediately.",
                    ],
                },
            ],
        },
        {
            "num": "5",
            "title": "Key Features",
            "bullets": [
                "Structured design canvas: 18 shape types covering common design blocks (buttons, headings, text, "
                "images, cards, lists, and more) with properties for labels, colors, and dimensions.",
                "Real-time collaboration: live object edits, presence, and cursor awareness via Liveblocks.",
                "Project hierarchy: teams own projects; projects contain pages; pages contain shapes.",
                "Role-based access control: owner / editor / viewer permissions enforced server-side.",
                "Invite and access-request workflows with real-time notifications.",
                "AI layout generation: local Ollama-powered assistant that returns structured, editable shapes.",
                "Public marketing site with a rendered system-architecture diagram.",
                "Swagger-documented REST API with a health-check endpoint.",
                "Retry-capable HTTP client and graceful fallback when the AI service is unavailable.",
            ],
        },
        {
            "num": "6",
            "title": "Technology Stack",
            "table": {
                "headers": ["Layer", "Technology", "Purpose"],
                "rows": [
                    ["Frontend", "React 19, Vite, TypeScript", "Application UI and build tooling"],
                    ["Canvas", "Konva / react-konva", "Rendering and interaction of design shapes"],
                    ["Real-time sync", "Liveblocks (React) + @liveblocks/node", "Rooms, presence, CRDT state, webhooks"],
                    ["Notifications", "Socket.IO (client + server)", "Real-time access events and updates"],
                    ["Backend", "Node.js, Express 5, TypeScript", "REST API, business logic, AI proxy"],
                    ["Data store", "PostgreSQL + Prisma ORM", "Persistent storage of users, teams, projects, pages, shapes"],
                    ["Authentication", "Clerk", "Sign-in, JWT issuance and verification"],
                    ["AI", "Express route + local Ollama (qwen2.5)", "Structured layout generation from prompts"],
                    ["Docs / API", "Swagger (swagger-ui-express)", "Interactive API reference"],
                ],
            },
        },
        {
            "num": "7",
            "title": "System Architecture",
            "paras": [
                "The system is organized into four logical layers. External cloud services (Liveblocks, Clerk) handle "
                "real-time sync and identity. The client is a React single-page application with a dashboard, a "
                "marketing landing page, and the canvas editor. The application server exposes the REST API, verifies "
                "JWTs, enforces role-based access, proxies AI requests to Ollama, and consumes Liveblocks webhooks. "
                "Data stores consist of PostgreSQL for durable records and the local Ollama model service for AI "
                "inference.",
                "Requests flow through a small set of well-defined surfaces: the API receives authenticated HTTP "
                "calls; Liveblocks pushes storageUpdated webhooks to the server; Socket.IO carries real-time "
                "notifications to clients; and Ollama serves model inference to the AI route. This separation keeps "
                "responsibilities clean and makes each layer independently testable and scalable.",
            ],
        },
        {
            "num": "8",
            "title": "Security & Access Control",
            "bullets": [
                "Authentication: all protected routes require a valid Clerk JWT, verified by a requireAuth "
                "middleware before any handler runs.",
                "Authorization: a requireProjectRole middleware enforces a three-level hierarchy \u2014 owner (3), "
                "editor (2), viewer (1) \u2014 and rejects insufficient-privilege requests with 403.",
                "Access requests: a request-access endpoint lets a user ask for access to a project; only owners can "
                "approve or deny, and the dashboard surfaces pending requests.",
                "Webhook security: Liveblocks webhook payloads are verified against the configured secret.",
                "CORS: the API restricts origins to the allowed client origins rather than accepting all requests.",
                "Input hygiene: AI prompt payloads are validated and length-clamped before being sent to Ollama.",
                "Credential management: secrets are supplied through environment variables and are not committed to "
                "the repository.",
            ],
        },
        {
            "num": "9",
            "title": "Testing & Quality Assurance",
            "paras": [
                "Verification covered the critical user journeys end to end, from a fresh sign-in through canvas "
                "editing, collaboration, AI generation, and access management. The table below records the test "
                "scenarios and their outcomes.",
            ],
            "table": {
                "headers": ["ID", "Area", "Scenario", "Expected Result", "Status"],
                "rows": [
                    ["T-01", "Auth", "User signs in with email via Clerk", "JWT issued; redirect to dashboard; profile shown", "Passed"],
                    ["T-02", "Auth", "Unauthenticated request to protected API route", "Request rejected with 401", "Passed"],
                    ["T-03", "API", "Create project", "Project persisted and returned with UUID and timestamps", "Passed"],
                    ["T-04", "API", "List projects for the signed-in user", "Paged list with search and pagination metadata", "Passed"],
                    ["T-05", "API", "Create a page within a project", "Page persisted and linked to project (200)", "Passed"],
                    ["T-06", "API", "List pages with pagination", "Page data returned with pagination metadata", "Passed"],
                    ["T-07", "API", "List project members with roles", "Members with role levels returned (200)", "Passed"],
                    ["T-08", "Canvas", "Add, move, resize, and delete shapes", "Canvas updates instantly and persists", "Passed"],
                    ["T-09", "Collaboration", "Two clients edit the same page simultaneously", "Changes synchronize live; presence visible", "Passed"],
                    ["T-10", "Collaboration", "Late joiner opens an existing page", "Current room state loaded immediately", "Passed"],
                    ["T-11", "AI", "Prompt \u201chero section\u2026\u201d", "Valid structured shape JSON returned and rendered", "Passed"],
                    ["T-12", "AI", "Ollama unavailable", "Graceful error with status surfaced to user", "Passed"],
                    ["T-13", "Access", "Request access to a project", "Pending request visible to project owner", "Passed"],
                    ["T-14", "Access", "Owner approves / denies a request", "Role updated or request rejected with notification", "Passed"],
                    ["T-15", "Security", "Viewer attempts edit-only action", "Request rejected with 403 by server", "Passed"],
                    ["T-16", "Build", "Server type-check (tsc)", "No type errors", "Passed"],
                    ["T-17", "Build", "Production client build (vite build)", "Build completes without errors", "Passed"],
                    ["T-18", "Site", "Landing page renders architecture diagram", "Diagram and sections render correctly", "Passed"],
                ],
            },
        },
        {
            "num": "10",
            "title": "Deployment & Operations",
            "bullets": [
                "Prerequisites: Node.js, PostgreSQL, an Ollama instance serving a chat model, and accounts for "
                "Liveblocks and Clerk.",
                "Configuration: all keys and database credentials are supplied via a .env file documented in the "
                "repository; no secrets are stored in source control.",
                "Local development: the API runs with ts-node in watch mode, the client runs through Vite with a "
                "proxy to the API, and the database schema is managed with Prisma migrations.",
                "Health monitoring: a /api/health endpoint reports server and database status for uptime checks.",
                "API documentation: a Swagger UI is served alongside the API for consumers and QA.",
            ],
        },
        {
            "num": "11",
            "title": "Performance & Scalability Considerations",
            "bullets": [
                "Persistence is asynchronous: webhook-based writes keep the database off the interactive editing path.",
                "The canvas relies on Konva's layered rendering, which comfortably handles hundreds of shape nodes.",
                "The editor is lazy-loaded so the dashboard and landing page stay light.",
                "The HTTP client retries transient failures, reducing spurious user-facing errors.",
                "AI inference is local, so generation cost is bounded by hardware and incurs no per-call cloud fees.",
                "Rooms are isolated per page, so collaboration load scales horizontally with concurrent editing sessions.",
            ],
        },
        {
            "num": "12",
            "title": "Roadmap",
            "bullets": [
                "Add caching (e.g., Redis) for high-read paths such as the dashboard and project listing.",
                "Introduce a work queue for webhook processing and notifications to absorb bursts under load.",
                "Extract a shared types package to keep client and server contracts synchronized.",
                "Extend the AI assistant to image generation and richer layout presets.",
                "Expand automated test coverage (unit and integration) to lock in current behavior.",
                "Add export of canvases to PNG/SVG and support for page templates.",
            ],
        },
        {
            "num": "Appendix A",
            "title": "API Endpoint Reference",
            "paras": [
                "The REST API is grouped by domain: authentication, projects, pages, members, access requests, "
                "notifications, and AI. Key routes include:",
            ],
            "table": {
                "headers": ["Method", "Route", "Description"],
                "rows": [
                    ["GET", "/api/health", "Server and database health check"],
                    ["GET/POST", "/api/projects", "List or create projects"],
                    ["GET/PATCH/DELETE", "/api/projects/:projectId", "Fetch, update, or delete a project"],
                    ["GET/POST", "/api/projects/:projectId/pages", "List or create pages"],
                    ["GET", "/api/projects/:projectId/members", "List project members with roles"],
                    ["POST", "/api/access-request", "Request access to a project"],
                    ["GET/POST", "/api/notifications", "List or create notifications"],
                    ["POST", "/api/ai/generate", "Generate structured layout from a prompt"],
                ],
            },
        },
    ],
}


def build_html():
    css = """
    * { box-sizing: border-box; }
    html, body { margin: 0; padding: 0; }
    body {
        font-family: 'Segoe UI', Calibri, Arial, sans-serif;
        color: #0f172a; line-height: 1.55; font-size: 11pt;
        background: #f1f5f9;
    }
    .page {
        background: #ffffff; max-width: 210mm; margin: 0 auto;
        padding: 18mm 20mm; box-shadow: 0 0 10px rgba(15,23,42,.12);
        min-height: 297mm;
    }
    .cover { text-align: center; padding-top: 70mm; }
    .cover .brand {
        display: inline-block; padding: 12px 28px; border-radius: 10px;
        background: linear-gradient(135deg, #0ea5e9, #0369a1); color: #fff;
        font-size: 34pt; font-weight: 700; letter-spacing: 1px;
    }
    .cover h1 { font-size: 22pt; margin: 14mm 0 2mm; color: #0f172a; }
    .cover .subtitle { font-size: 13pt; color: #475569; }
    .cover .rule { width: 60mm; height: 3px; margin: 8mm auto; background: #0ea5e9; }
    .cover .meta { font-size: 10.5pt; color: #475569; }
    h2 {
        font-size: 15pt; color: #0369a1; border-bottom: 2px solid #0ea5e9;
        padding-bottom: 3px; margin: 14pt 0 8pt;
    }
    h3 { font-size: 12.5pt; color: #0f172a; margin: 10pt 0 5pt; }
    p { margin: 5pt 0; text-align: justify; }
    ul { margin: 4pt 0 6pt 16pt; padding: 0; }
    li { margin: 2.5pt 0; }
    table { border-collapse: collapse; width: 100%; margin: 8pt 0; font-size: 9pt; }
    th {
        background: #0ea5e9; color: #fff; text-align: left;
        padding: 6px 8px; border: 1px solid #0284c7; font-weight: 600;
    }
    td { padding: 5px 8px; border: 1px solid #cbd5e1; vertical-align: top; }
    tr:nth-child(even) td { background: #f8fafc; }
    .status { color: #047857; font-weight: 600; }
    .footer {
        margin-top: 20pt; padding-top: 6pt; border-top: 1px solid #cbd5e1;
        font-size: 8.5pt; color: #94a3b8; display: flex; justify-content: space-between;
    }
    @page { size: A4; margin: 0; }
    @media print {
        body { background: #fff; }
        .page { box-shadow: none; margin: 0; max-width: none; min-height: 0; page-break-after: always; }
    }
    """
    parts = ["<!DOCTYPE html><html><head><meta charset='utf-8'>",
             "<title>Craftboard Project Overview</title><style>", css, "</style></head><body>"]

    parts.append("<div class='page'>")
    parts.append("<div class='cover'>")
    parts.append("<div class='brand'>%s</div>" % html.escape(TITLE))
    parts.append("<h1>%s</h1>" % html.escape(SUBTITLE))
    parts.append("<div class='rule'></div>")
    parts.append("<div class='subtitle'>%s</div>" % html.escape(DOC_TITLE))
    parts.append("<div class='meta'><br>Version %s &nbsp;|&nbsp; %s<br>Prepared by: Development Team</div>"
                 % (VERSION, TODAY))
    parts.append("</div></div>")

    for sec in CONTENT["sections"]:
        parts.append("<div class='page'>")
        parts.append("<h2>%s &nbsp; %s</h2>" % (html.escape(sec["num"]), html.escape(sec["title"])))
        for p in sec.get("paras", []):
            parts.append("<p>%s</p>" % html.escape(p))
        for b in sec.get("bullets", []):
            parts.append("<ul><li>%s</li></ul>" % html.escape(b))
        for sub in sec.get("subsections", []):
            parts.append("<h3>%s</h3>" % html.escape(sub["title"]))
            for p in sub.get("paras", []):
                parts.append("<p>%s</p>" % html.escape(p))
        for b in sec.get("paras2", []):
            parts.append("<p>%s</p>" % html.escape(b))
        if "table" in sec:
            tbl = sec["table"]
            parts.append("<table><thead><tr>")
            for h in tbl["headers"]:
                parts.append("<th>%s</th>" % html.escape(h))
            parts.append("</tr></thead><tbody>")
            for row in tbl["rows"]:
                parts.append("<tr>")
                for i, cell in enumerate(row):
                    cls = "status" if (tbl["headers"][-1] == "Status" and i == len(row) - 1) else ""
                    parts.append("<td class='%s'>%s</td>" % (cls, html.escape(cell)))
                parts.append("</tr>")
            parts.append("</tbody></table>")
        parts.append("<div class='footer'><span>Craftboard &mdash; Project Overview &amp; Technical Documentation</span>"
                     "<span>Version %s &bull; %s</span></div>" % (VERSION, TODAY))
        parts.append("</div>")

    parts.append("</body></html>")
    return "".join(parts)


def build_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        st = doc.styles[style_name]
        st.font.name = "Calibri"
        st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    for hname, size, color in (("Title", 34, "0369A1"), ("Heading 1", 16, "0369A1"),
                               ("Heading 2", 13, "0F172A")):
        st = doc.styles[hname]
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = True

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(TITLE)
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0369A1")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(SUBTITLE)
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    docl = doc.add_paragraph()
    docl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = docl.add_run(DOC_TITLE)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Version %s  |  %s\nPrepared by: Development Team" % (VERSION, TODAY))
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_page_break()

    for sec in CONTENT["sections"]:
        h = doc.add_heading("%s  %s" % (sec["num"], sec["title"]), level=1)
        for p in sec.get("paras", []):
            doc.add_paragraph(p)
        for b in sec.get("bullets", []):
            doc.add_paragraph(b, style="List Bullet")
        for sub in sec.get("subsections", []):
            doc.add_heading(sub["title"], level=2)
            for p in sub.get("paras", []):
                doc.add_paragraph(p)
        for b in sec.get("paras2", []):
            doc.add_paragraph(b)
        if "table" in sec:
            tbl = sec["table"]
            t = doc.add_table(rows=1, cols=len(tbl["headers"]))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = t.rows[0].cells
            for i, htext in enumerate(tbl["headers"]):
                hdr[i].paragraphs[0].add_run(htext).bold = True
            for row in tbl["rows"]:
                cells = t.add_row().cells
                for i, cell in enumerate(row):
                    cells[i].text = cell

    doc.save(DOCX_PATH)


def main():
    with open(HTML_PATH, "w", encoding="utf-8") as fh:
        fh.write(build_html())
    build_docx()
    print("Wrote:", HTML_PATH)
    print("Wrote:", DOCX_PATH)


if __name__ == "__main__":
    main()
